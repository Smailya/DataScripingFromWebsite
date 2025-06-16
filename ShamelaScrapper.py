import requests
from bs4 import BeautifulSoup
from docx import Document
from deep_translator import GoogleTranslator
import time

# Configuration
START_PAGE = 119
END_PAGE = 126
BASE_URL = "https://shamela.ws/book/96463/{page}#p1"
OUTPUT_FILE = "shamela_translated_pages.docx"

# Initialize components
doc = Document()
doc.add_heading('Shamela.ws Content - Arabic & English', 0)
translator = GoogleTranslator(source='auto', target='en')

headers = {
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/58.0.3029.110 Safari/537.3"
}

def translate_text(text):
    try:
        return translator.translate(text)
    except Exception as e:
        print(f"Translation error: {e}")
        return "Translation unavailable"

for page_num in range(START_PAGE, END_PAGE + 1):
    url = BASE_URL.format(page=page_num)
    print(f"Processing page {page_num}...")
    
    try:
        response = requests.get(url, headers=headers, timeout=15)
        response.raise_for_status()
    except Exception as e:
        print(f"Page {page_num} fetch failed: {e}")
        continue

    soup = BeautifulSoup(response.text, "html.parser")
    content_div = soup.find("div", class_="nass")
    
    if not content_div:
        print(f"No content found on page {page_num}")
        continue

    # Extract and clean Arabic text
    arabic_text = content_div.get_text(separator="\n", strip=True)
    
    # Translate to English
    english_text = translate_text(arabic_text)
    
    # Add to document
    doc.add_heading(f'Page {page_num}', level=1)
    doc.add_paragraph("Arabic Text:\n" + arabic_text)
    doc.add_paragraph("\nEnglish Translation:\n" + english_text)
    
    time.sleep(2)  # Increased delay for translation compliance

# Save document
doc.save(OUTPUT_FILE)
print(f"Successfully saved bilingual document to {OUTPUT_FILE}")
